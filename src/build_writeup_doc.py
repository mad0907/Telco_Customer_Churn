"""
Builds the Word write-up: docs/Telco_Churn_Writeup_and_Architecture.docx
Combines the analysis write-up with the full deployment/architecture
blueprint and AI governance section. Numbers embedded below were pulled
directly from the executed notebook (notebooks/telco_churn_analysis.ipynb)
to keep the two documents consistent.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEXT_COLOR = RGBColor(0x10, 0x18, 0x28)
MUTED_COLOR = RGBColor(0x47, 0x54, 0x69)
ACCENT_COLOR = RGBColor(0x1D, 0x4E, 0xD8)
RULE_COLOR = "344054"

doc = Document()

# --- base styles -----------------------------------------------------------
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.font.color.rgb = TEXT_COLOR

section = doc.sections[0]
section.page_width = Cm(21.59)   # US Letter
section.page_height = Cm(27.94)
section.left_margin = Cm(2.2)
section.right_margin = Cm(2.2)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)


def add_heading(text, level=1, space_before=18, space_after=6):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.color.rgb = TEXT_COLOR
    h.paragraph_format.space_before = Pt(space_before)
    h.paragraph_format.space_after = Pt(space_after)
    return h


def add_para(text="", bold=False, italic=False, size=11, color=None, space_after=8, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color or TEXT_COLOR
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    return p


def add_rich_para(segments, size=11, space_after=8):
    """segments: list of (text, bold, italic) tuples"""
    p = doc.add_paragraph()
    for text, bold, italic in segments:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.color.rgb = TEXT_COLOR
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullets(items, size=11, space_after=4):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        if isinstance(item, tuple):
            lead, rest = item
            r1 = p.add_run(lead)
            r1.bold = True
            r1.font.size = Pt(size)
            r1.font.color.rgb = TEXT_COLOR
            r2 = p.add_run(rest)
            r2.font.size = Pt(size)
            r2.font.color.rgb = TEXT_COLOR
        else:
            r = p.add_run(item)
            r.font.size = Pt(size)
            r.font.color.rgb = TEXT_COLOR
        p.paragraph_format.space_after = Pt(space_after)


def set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_table(headers, rows, col_widths_cm, header_fill="344054", size=9.5):
    """Add a styled table with a colored header row and zebra-striped body rows."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, w in enumerate(col_widths_cm):
        table.columns[i].width = Cm(w)

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].width = Cm(col_widths_cm[i])
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr_cells[i], header_fill)

    for r_i, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].width = Cm(col_widths_cm[i])
            p = cells[i].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(size)
            run.font.color.rgb = TEXT_COLOR
            if r_i % 2 == 1:
                set_cell_shading(cells[i], "F2F4F7")
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_code_block(code_text, size=8.5):
    """Add a monospace, shaded, bordered paragraph for a code/config excerpt."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F2F4F7")
    pPr.append(shd)
    border = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "D0D5DD")
        border.append(el)
    pPr.append(border)
    for line in code_text.strip("\n").split("\n"):
        run = p.add_run(line + "\n")
        run.font.name = "Consolas"
        run.font.size = Pt(size)
        run.font.color.rgb = TEXT_COLOR


def add_image(path, width_in=6.4, caption=None):
    """Insert a centered image with an optional italic caption below it."""
    doc.add_picture(path, width=Inches(width_in))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cap = add_para(caption, italic=True, size=9, color=MUTED_COLOR, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
        return cap


def add_rule():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), RULE_COLOR)
    border.append(bottom)
    pPr.append(border)


def add_page_break():
    doc.add_page_break()


# ===========================================================================
# TITLE PAGE
# ===========================================================================
doc.add_paragraph().paragraph_format.space_after = Pt(60)
add_para("Telco Customer Churn", bold=True, size=28, space_after=4)
add_para("Predictive Modeling, Deployment Architecture, and Governance", size=16, color=MUTED_COLOR, space_after=40)
add_rule()
add_para("")
add_rich_para([("Prepared by: ", True, False), ("Munmun", False, False)], size=12, space_after=4)
add_rich_para([("Date: ", True, False), ("August 2026", False, False)], size=12, space_after=4)
add_rich_para([("Dataset: ", True, False), ("IBM/Kaggle Telco Customer Churn", False, False)], size=12, space_after=4)
add_para("")
add_para(
    "This document accompanies the analysis notebook "
    "(notebooks/telco_churn_analysis.ipynb) and the deployment blueprint "
    "(deployment/) delivered alongside it. It does not restate every step "
    "shown in the notebook -- it summarizes results, states the reasoning "
    "behind the choices made, and covers the deployment, governance, and "
    "monitoring design that sits outside the notebook itself.",
    italic=True, size=10.5, color=MUTED_COLOR,
)
add_page_break()

# ===========================================================================
# TABLE OF CONTENTS (manual, since this is a fixed deliverable)
# ===========================================================================
add_heading("Contents", level=1)
toc_items = [
    "Part I -- Analysis Write-Up",
    "1. Executive Summary",
    "2. Assumptions",
    "3. Methodology Summary",
    "4. Results",
    "5. Model Selection and Recommendation",
    "6. Model Interpretation -- Drivers of Churn",
    "7. Post-Deployment Monitoring",
    "8. Limitations and Next Steps",
    "Part II -- Deployment Architecture and Governance Blueprint",
    "9. Solution Architecture",
    "10. MLOps Pipeline -- Retraining and Promotion",
    "11. Serving Layer -- Managed Online Endpoint",
    "12. AI Governance and Accountability",
    "13. Immutable Prediction Logging",
    "14. Continuous Improvement Strategy",
    "15. Marketing Mix Model -- Applicability Assessment",
    "16. Azure Platform Capabilities Referenced",
    "17. Deliverable Manifest",
]
for item in toc_items:
    is_part = item.startswith("Part")
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.bold = is_part
    run.font.size = Pt(12.5 if is_part else 11)
    run.font.color.rgb = ACCENT_COLOR if is_part else TEXT_COLOR
    p.paragraph_format.space_before = Pt(10 if is_part else 0)
    p.paragraph_format.space_after = Pt(4 if is_part else 3)
add_page_break()

# ===========================================================================
# PART I -- ANALYSIS WRITE-UP
# ===========================================================================
add_heading("Part I -- Analysis Write-Up", level=1)

add_heading("1. Executive Summary", level=2)
add_para(
    "This document evaluates whether Telco customer churn can be predicted "
    "from account and service data, which factors drive that risk, and how "
    "a model built for this purpose should be deployed, monitored, and "
    "governed once it is making or informing real retention decisions."
)
add_bullets([
    ("Approach: ", "three models were built and compared -- a Logistic "
     "Regression baseline, a Random Forest (bagging), and an XGBoost "
     "(boosting) model -- evaluated with threshold-independent metrics "
     "(ROC-AUC, PR-AUC) and a cost-sensitive threshold analysis, not "
     "accuracy alone."),
    ("Result: ", "all three models perform comparably (ROC-AUC 0.845-0.847); "
     "a bootstrap analysis confirmed the cost differences between them are "
     "not statistically distinguishable on this dataset. XGBoost is "
     "recommended for deployment on secondary grounds -- its native fit "
     "with the interaction effects found in EDA and with SHAP-based "
     "explainability at serving time -- not because it wins outright on "
     "a single metric."),
    ("Drivers: ", "contract type, tenure, internet service type, and the "
     "presence or absence of protective add-ons (tech support, online "
     "security) are the dominant churn drivers, confirmed independently "
     "through mutual information ranking and SHAP."),
    ("Deployment: ", "Part II of this document specifies how the "
     "recommended model would be deployed as a governed Azure ML service, "
     "including the MLOps retraining pipeline, an immutable audit-logging "
     "design for every prediction, and the guardrails needed before this "
     "model is allowed to influence an automated retention decision."),
])

add_heading("2. Assumptions", level=2)
add_bullets([
    "The dataset is a single snapshot, not a longitudinal panel. Tenure at "
    "the time of the snapshot is used as a proxy for time-to-churn; the "
    "survival-curve analysis in the notebook should be read as directionally "
    "correct rather than a precise hazard estimate a true panel dataset "
    "would support.",
    "Cost figures used for threshold selection ($400 per missed churner, "
    "$50 per unnecessary retention offer) are illustrative, stated "
    "assumptions -- the dataset contains no acquisition-cost or "
    "campaign-cost fields. These should be replaced with Finance-provided "
    "figures before the recommended operating threshold is used in "
    "production.",
    "No marketing-spend or campaign data exists in this dataset. This "
    "directly affects the scope of Section 15 -- a Marketing Mix Model was "
    "considered and found not applicable to this dataset for that reason.",
    "\"Development done locally\" (per the assignment brief) is interpreted "
    "as executed in a local Python virtual environment with saved outputs, "
    "delivered as a notebook that can additionally be opened in a hosted "
    "notebook environment for ease of review.",
    "This dataset is a single cross-sectional snapshot with no defined "
    "observation/prediction/outcome window -- no field is generated after "
    "a churn decision (no leakage in that specific sense), but a production "
    "system would need that temporal structure explicitly, and would use a "
    "time-based train/test split rather than the random stratified split "
    "used here, once timestamped data exists. See the notebook's Section 1 "
    "leakage assessment for the full reasoning.",
])

add_heading("3. Methodology Summary", level=2)
add_para(
    "Full detail, code, and narrative justification for every step below "
    "is in notebooks/telco_churn_analysis.ipynb. This section summarizes "
    "what was done and why; it is not a substitute for the notebook."
)
add_bullets([
    ("Data cleaning: ", "the 11 blank TotalCharges values were identified "
     "as a structural zero (all have tenure = 0, i.e. day-one customers with "
     "no completed billing cycle) and set to 0.0, rather than imputed with "
     "a median or dropped. An IQR outlier scan on tenure, MonthlyCharges, "
     "and TotalCharges found no outliers -- explicitly checked, not "
     "assumed."),
    ("EDA beyond the standard view: ", "a Kaplan-Meier survival curve showed "
     "churn hazard is front-loaded in the first several months rather than "
     "uniform across tenure; a revenue-at-risk analysis found churned "
     "customers are, on average, higher-value than retained customers "
     "($893/year vs. $735/year), not lower-value as might be assumed; "
     "a mutual-information ranking (not Pearson correlation, which is the "
     "wrong tool for nominal categories) identified Contract, "
     "OnlineSecurity, TechSupport, and InternetService as the strongest "
     "categorical predictors; and a bundling analysis showed churn declines "
     "close to monotonically with the number of add-on services held."),
    ("Feature engineering: ", "a service-count feature, an implied "
     "average-monthly-spend feature, and a tenure-bucket feature were "
     "derived, each justified against a specific EDA finding rather than "
     "added speculatively."),
    ("Class imbalance: ", "handled via class_weight/scale_pos_weight inside "
     "each estimator rather than synthetic oversampling -- with 1,869 real "
     "churn examples in the dataset, resampling was judged unnecessary and "
     "would add synthetic-data risk without a clear benefit."),
])

add_heading("4. Results", level=2)
add_para("Threshold-independent performance on the held-out test set:")
add_table(
    headers=["Model", "ROC-AUC", "PR-AUC", "Precision@0.5", "Recall@0.5", "F1@0.5"],
    rows=[
        ["Logistic Regression", "0.845", "0.652", "0.505", "0.802", "0.620"],
        ["Random Forest", "0.846", "0.662", "0.527", "0.797", "0.634"],
        ["XGBoost", "0.847", "0.659", "0.517", "0.805", "0.630"],
    ],
    col_widths_cm=[4.0, 2.2, 2.2, 2.4, 2.2, 2.0],
)
add_para(
    "At the cost-minimizing threshold from the cost-sensitive analysis "
    "(illustrative cost assumptions: $400 false negative, $50 false "
    "positive):"
)
add_table(
    headers=["Model", "Threshold", "Precision", "Recall", "F1", "Median Cost ($)", "95% CI ($)"],
    rows=[
        ["Logistic Regression", "0.27", "0.423", "0.936", "0.582", "33,500", "29,648 - 37,950"],
        ["Random Forest", "0.22", "0.389", "0.960", "0.554", "34,200", "30,998 - 37,650"],
        ["XGBoost", "0.30", "0.425", "0.930", "0.583", "33,850", "29,800 - 38,450"],
    ],
    col_widths_cm=[3.6, 2.0, 2.0, 1.8, 1.6, 2.6, 3.0],
    size=8.5,
)
add_para(
    "The bootstrapped 95% confidence intervals overlap substantially across "
    "all three models. That overlap is itself the key result of this table: "
    "there is no statistically defensible cost winner among the three "
    "models on this dataset and this cost assumption -- a conclusion that "
    "would have been missed by reporting the point estimates alone.",
    italic=True, color=MUTED_COLOR, size=10,
)
add_para(
    "The cost-sensitive threshold assumes the retention team can act on "
    "every flagged customer. In practice, contact capacity is usually "
    "fixed, so the notebook (Section 9.4) also evaluates a capacity-"
    "constrained view: all three models capture close to 28% of actual "
    "churners within the top 10% of customers ranked by risk -- a lift of "
    "roughly 2.8x over random targeting. This is the number to hand to a "
    "team with a fixed weekly contact capacity; it is a distinct question "
    "from what the cost-sensitive threshold above answers."
)
add_para(
    "Model output was also translated into a risk-tier scheme (Low / "
    "Medium / High / Critical), validated against observed test-set churn "
    "rate rather than left as an arbitrary label:"
)
add_table(
    headers=["Tier", "Customers", "Observed Churn Rate"],
    rows=[
        ["Low", "590", "4.4%"],
        ["Medium", "363", "23.7%"],
        ["High", "291", "48.8%"],
        ["Critical", "165", "72.7%"],
    ],
    col_widths_cm=[3.0, 3.0, 4.0],
)
add_para(
    "Churn rate rises monotonically from Low to Critical, which is the "
    "check that matters -- a tier scheme where \"Critical\" customers do "
    "not actually churn more often than \"Low\" customers would be a broken "
    "segmentation regardless of how intuitive the cut points look. Note "
    "that even the Critical tier is 27.3% non-churners, which is why the "
    "recommended action for that tier (Section 6) is a retention offer, "
    "not a certainty-based action."
)
add_para(
    "Finally, ranking customers by churn probability alone is not the same "
    "as ranking by revenue protected. Weighting each customer's probability "
    "by their projected annual value (probability x MonthlyCharges x 12) "
    "and comparing the resulting top 10 against the top 10 by raw "
    "probability found only 1 customer in common between the two lists -- "
    "a far larger divergence than intuition would suggest. For a capacity-"
    "constrained retention team, the value-weighted ranking is the more "
    "defensible prioritization metric, since it directly answers \"where "
    "does intervention protect the most revenue\" rather than only \"who is "
    "statistically most likely to leave.\""
)
add_page_break()

add_heading("5. Model Selection and Recommendation", level=2)
add_rich_para([("Recommendation: ", True, False), ("XGBoost", True, False),
               (", operated at its cost-minimizing threshold (0.30), with "
                "Logistic Regression and Random Forest retained as "
                "documented, statistically-tied alternatives.", False, False)])
add_para(
    "Because the bootstrap analysis rules out choosing a winner on cost "
    "alone, the recommendation rests on criteria the point estimate cannot "
    "capture:"
)
add_bullets([
    ("Robustness to interaction effects. ", "Churn risk in this data depends "
     "on combinations of features -- bundled services, fiber internet "
     "without protective add-ons -- not single variables in isolation. "
     "XGBoost and Random Forest capture these natively; Logistic Regression "
     "would need every such interaction hand-engineered to compete, adding "
     "maintenance burden as new interactions are discovered."),
    ("Interpretability infrastructure fit. ", "XGBoost pairs directly with "
     "SHAP's efficient tree explainer, fast enough to support the "
     "confidence-based human-review routing specified in the serving layer "
     "(Section 11). Random Forest shares this property but explains more "
     "slowly at comparable tree counts."),
    ("Operational fit. ", "XGBoost's training and scoring cost, and its "
     "compatibility with Azure ML managed online endpoints and the "
     "champion/challenger promotion gate (Section 10), make it a practical "
     "default for the retraining cadence in Section 14."),
])
add_para(
    "Logistic Regression is not discarded: it remains the standing "
    "baseline, the fastest model to retrain for a sanity check, and the "
    "most directly auditable model if a regulator or internal model-risk "
    "review needs a coefficient-level explanation rather than a SHAP "
    "approximation. This recommendation should be re-run once real "
    "Finance-provided cost figures replace the illustrative assumption in "
    "Section 4."
)

add_heading("6. Model Interpretation -- Drivers of Churn", level=2)
add_para(
    "SHAP values on the recommended XGBoost model were used to decompose "
    "predictions into per-feature contributions -- this answers \"what "
    "drives churn and in which direction,\" which a raw feature-importance "
    "ranking does not. The ranking independently agrees with the "
    "mutual-information ranking computed directly from the raw data in "
    "EDA, which is a useful cross-check rather than a coincidence."
)
add_bullets([
    ("Contract type: ", "month-to-month contracts push predictions toward "
     "churn; one- and two-year contracts push firmly the other way -- "
     "consistent with the survival curve showing hazard concentrated in "
     "early tenure, which month-to-month customers are exposed to every "
     "renewal cycle."),
    ("Tenure: ", "low tenure pushes toward churn, consistent with the "
     "survival analysis."),
    ("Internet service and add-ons: ", "fiber-optic customers without "
     "security or tech-support add-ons show materially elevated churn risk "
     "compared to DSL customers or fiber customers with those add-ons -- an "
     "interaction effect, not a standalone service effect."),
    ("Payment method: ", "electronic-check payment pushes toward churn, "
     "persisting across paperless-billing status -- read as a proxy for a "
     "less-engaged, more price-sensitive customer segment rather than a "
     "friction effect of paperless billing itself."),
])
add_para(
    "A SHAP summary plot is an analyst's view, not a retention agent's. "
    "The notebook (Section 11.1) also includes a small function that turns "
    "a single prediction into a decision card -- risk tier, the specific "
    "drivers behind that customer's score, and a recommended action tied "
    "to those drivers, not a generic one. The same logic is exposed "
    "interactively in a small prototype delivered alongside this analysis: "
    "app.py, a Streamlit application that lets a user enter a customer "
    "profile and see the prediction, risk tier, top SHAP drivers, and "
    "recommended action in real time, loading the exact model persisted by "
    "the notebook rather than retraining. It is a decision-support "
    "prototype, not the production serving surface -- the production "
    "equivalent of the same logic is deployment/score.py behind the "
    "managed online endpoint described in Part II. Setup and run "
    "instructions are in README.md."
)

add_heading("7. Post-Deployment Monitoring", level=2)
add_para(
    "A model that scores well today degrades silently in production for "
    "two distinct reasons, each needing its own monitor:"
)
add_bullets([
    ("Data drift", " -- the distribution of incoming records diverges from "
     "training (a new pricing plan, a new payment provider). Tracked via "
     "Population Stability Index (PSI) per feature, computed daily "
     "(deployment/monitoring_job.yml) against the training-set reference "
     "distribution. A PSI above 0.2 on a top-ranked feature (Contract, "
     "tenure, InternetService, PaymentMethod) triggers investigation before "
     "it triggers a retrain."),
    ("Concept drift", " -- the relationship between features and churn "
     "itself shifts (competitor pricing changes what \"expensive\" means to "
     "customers). PSI cannot detect this since input distributions can stay "
     "stable while the true relationship moves. Tracked via rolling "
     "7/30-day recall and precision once true outcomes are observed and "
     "joined back to historical predictions, compared to the Section 4 "
     "baseline. A sustained drop, not a single noisy week, triggers "
     "retraining."),
    ("Operational monitors", " -- prediction-volume and class-rate drift "
     "(often a pipeline issue before a model issue), endpoint latency and "
     "error rate, and feature-level null-rate monitoring (a silent upstream "
     "schema change would not be caught by PSI if imputation quietly "
     "absorbs it)."),
])
add_para(
    "Retraining cadence: quarterly as a floor, with an unscheduled retrain "
    "triggered immediately by a PSI breach on a top-ranked feature or a "
    "sustained recall/precision drop -- whichever comes first. Every "
    "retrained candidate is evaluated against the current production model "
    "on the cost-sensitive metric before promotion (champion/challenger), "
    "never deployed automatically on training completion. Full "
    "implementation detail is in Part II."
)

add_heading("8. Limitations and Next Steps", level=2)
add_para("Limitations:")
add_bullets([
    "Single-snapshot data limits how confidently the survival analysis "
    "generalizes; a true panel dataset with observed churn dates would "
    "support a full Cox proportional hazards model instead of a "
    "snapshot-based approximation.",
    "All three models cluster closely on every metric, including the "
    "bootstrapped cost estimate. The larger remaining lever is more likely "
    "additional features (support-ticket history, usage trend, competitor "
    "pricing exposure) than a different algorithm.",
])
add_para("Next steps:")
add_bullets([
    ("Uplift modeling. ", "Rather than only predicting who will churn, "
     "estimate the incremental effect of a specific retention action per "
     "customer (a T-learner or causal forest against historical campaign "
     "data, once such data exists). This is the correct tool for \"where "
     "does retention spend move the needle\" -- see Section 15 for why this "
     "supersedes a Marketing Mix Model for this use case."),
    ("A/B test", " the recommended model's retention-targeting policy "
     "against the current process before full rollout, rather than "
     "deploying on offline metrics alone."),
    ("Revisit the cost matrix", " in Section 4 with actual Finance-provided "
     "acquisition-cost and offer-cost figures, and rerun the bootstrap "
     "threshold analysis before committing to a production operating "
     "point."),
    ("Extend monitoring", " with the governance and audit-logging design in "
     "Part II before this model is exposed to any automated retention "
     "decision."),
])
add_page_break()

# ===========================================================================
# PART II -- DEPLOYMENT ARCHITECTURE AND GOVERNANCE BLUEPRINT
# ===========================================================================
add_heading("Part II -- Deployment Architecture and Governance Blueprint", level=1)
add_para(
    "This part goes beyond what the assignment strictly requires, to show "
    "how the recommended model would actually be operated as a governed "
    "service on Azure -- the environment this role is centered on. Full, "
    "runnable configuration files referenced below are delivered in the "
    "deployment/ folder alongside this document; this section explains the "
    "design decisions behind them rather than reproducing every file "
    "in full.",
    italic=True, color=MUTED_COLOR, size=10,
)

add_heading("9. Solution Architecture", level=2)
add_para(
    "The model is deployed as a scoring service behind a stable API, with "
    "training, serving, and governance kept as separate, independently "
    "scalable concerns rather than bundled into one pipeline."
)
add_image("../assets/architecture_diagram.png", width_in=6.5)
add_bullets([
    ("Ingestion. ", "Azure Data Factory pulls from CRM/billing, support "
     "tickets, and usage telemetry, validating schema before anything "
     "reaches the feature store -- the same validate-before-clean "
     "discipline used in the notebook's data audit, applied at the "
     "pipeline level."),
    ("Feature store. ", "Azure ML's managed feature store, so training and "
     "serving compute the same features from the same definitions -- the "
     "most common source of train/serve skew is a feature computed one way "
     "in a notebook and a subtly different way in the serving code, and a "
     "shared feature store removes that gap by construction."),
    ("Training and registry. ", "the training pipeline (Section 10) writes "
     "candidate models to the registry only after a champion/challenger "
     "comparison against the current production model; promotion to "
     "production is a separate, gated step."),
    ("Serving. ", "a managed online endpoint fronts the scoring API; the "
     "application/UI layer (the retention workflow) consumes predictions "
     "from it, never calling the model directly."),
    ("Governance layer. ", "every request/response pair is logged through "
     "the immutable-logging path (Section 13) before the response reaches "
     "the caller, and Azure Monitor / Application Insights feed the "
     "drift-detection job that can trigger an unscheduled retrain."),
])

add_heading("10. MLOps Pipeline -- Retraining and Promotion", level=2)
add_image("../assets/mlops_pipeline_diagram.png", width_in=6.5)
add_para(
    "The pipeline is deliberately structured so a retrained model cannot "
    "reach production without passing two independent gates: an automated "
    "one (does it beat the current champion on the cost-sensitive metric) "
    "and a human one (does a release manager sign off). Automating the "
    "comparison and keeping the sign-off manual is a deliberate choice -- "
    "the comparison is objective and should not depend on a person "
    "remembering to run it, but promotion to production is a business risk "
    "decision that should not be fully automated."
)
add_para("Training pipeline (deployment/training_pipeline.yml) -- key structure:")
add_code_block("""
jobs:
  validate_data: ...
  engineer_features: ...        # writes to the managed feature store
  train_candidate: ...          # XGBoost, cost-aware evaluation
  champion_challenger_compare:  # compares vs. current production model
    inputs:
      challenger_model: ${{parent.jobs.train_candidate.outputs.model}}
      production_model_name: telco-churn-xgboost
      production_model_stage: production
    outputs:
      decision: ...              # {"promote": true/false, "reason": "..."}
  register_if_better: ...        # only registers if decision.promote == true

schedule:
  frequency: week
  interval: 1
""")
add_para("CI/CD pipeline (deployment/azure-pipelines.yml) -- stage sequence:")
add_code_block("""
CI (lint, unit test, data validation)
  -> TrainAndEvaluate (submits training_pipeline.yml)
    -> ApprovalGate (Azure DevOps environment, required reviewers)
      -> CanaryDeploy (10% traffic to the new model version)
        -> HealthCheck (monitor error rate / latency for 30 minutes)
          -> FullRollout (100% traffic, only if the canary is healthy)
""")
add_para(
    "The canary stage exists because the champion/challenger comparison in "
    "the training pipeline is evaluated on historical test data -- it "
    "cannot catch a production-only failure mode (a missing environment "
    "variable, a serving-time schema mismatch) that only shows up under "
    "real traffic. A canary at 10% traffic limits the blast radius of "
    "exactly that class of failure."
)
add_page_break()

add_heading("11. Serving Layer -- Managed Online Endpoint", level=2)
add_para(
    "The endpoint and its deployment are defined separately "
    "(deployment/managed_online_endpoint.yml and "
    "managed_online_deployment.yml) so the public-facing endpoint name and "
    "auth stay stable across model versions, while individual deployments "
    "underneath it carry the traffic split used for canary rollout."
)
add_para(
    "The scoring entry script (deployment/score.py) does more than call "
    "predict() -- it is where the guardrails from Section 12 are actually "
    "enforced, not just described:"
)
add_bullets([
    ("Input validation before scoring. ", "a malformed or out-of-range "
     "request (negative tenure, a missing required field) is rejected "
     "before it reaches the model, rather than producing an unpredictable "
     "prediction from bad input."),
    ("Confidence-based review routing. ", "predictions within a defined "
     "margin of the operating threshold are flagged needs_human_review, so "
     "the application layer can route borderline cases to a person instead "
     "of an automated action -- an explicit accountability guardrail, not "
     "just a UI nicety."),
    ("Fail-closed audit logging. ", "if the immutable-logging call in "
     "Section 13 fails, the request fails -- the service never returns a "
     "prediction it could not log. This is a deliberate choice: a brief "
     "scoring outage during a logging incident is judged preferable to a "
     "silent gap in the audit trail."),
])
add_code_block("""
def _log_prediction_event(input_hash, prediction, probability):
    event = {"input_hash": input_hash, "model_version": MODEL_VERSION,
              "prediction": prediction, "probability": round(probability, 6),
              "timestamp_utc": datetime.now(timezone.utc).isoformat()}
    try:
        append_to_ledger(event)   # signed via Key Vault, written to
                                    # Azure Confidential Ledger -- see Section 13
    except Exception as exc:
        raise RuntimeError(
            "Prediction could not be recorded to the audit trail; "
            "refusing to return an unlogged prediction"
        ) from exc
""")
add_page_break()

add_heading("12. AI Governance and Accountability", level=2)
add_para(
    "A churn model that influences who receives a retention offer is a "
    "model risk, not just an analytics artifact, once it is deployed. The "
    "governance design below follows that framing."
)

add_heading("Model risk management framework", level=3)
add_bullets([
    "Every model version is documented with a model card: training data "
    "snapshot, features used, evaluation metrics, known limitations "
    "(Section 8), and the cost assumptions behind its operating threshold.",
    "No model reaches production without passing the pre-deployment "
    "validation gate: the champion/challenger comparison (Section 10) plus "
    "the bias and fairness check below.",
    "Every production promotion is signed off by a named release manager "
    "(the Azure DevOps approval-gate environment in Section 10), so a "
    "specific deployment can always be traced to a specific accountable "
    "person, not just a pipeline run.",
])

add_heading("Bias and fairness testing", level=3)
add_para(
    "This dataset includes SeniorCitizen and gender -- fields that require "
    "explicit disparate-impact testing before this model is allowed near a "
    "pricing or retention decision, not just a performance check. The "
    "recommended pre-deployment gate: compute recall, precision, and the "
    "predicted-positive rate separately for each SeniorCitizen and gender "
    "subgroup, and flag for review if any subgroup's predicted-positive "
    "rate differs by more than a defined tolerance (a standard disparate-"
    "impact check) from the overall population. If a gap is found, the "
    "response is not necessarily to drop the field outright -- SeniorCitizen "
    "in particular may carry genuine, non-discriminatory signal about "
    "service needs -- but the gap must be explained and documented before "
    "deployment, not discovered afterward from a complaint or an audit."
)

add_heading("Guardrails", level=3)
add_bullets([
    ("Input schema validation ", "at the API boundary (Section 11) -- "
     "malformed requests never reach the model."),
    ("Confidence-based human review ", "for predictions near the operating "
     "threshold, rather than fully automated action on every prediction "
     "regardless of certainty."),
    ("Fail-closed behavior ", "on both inference failure and audit-logging "
     "failure -- the system prefers no answer over an unaccountable one."),
    ("Explainability on demand, ", "not on every request: SHAP explanations "
     "are computed in the batch monitoring job and cached, with a separate, "
     "lower-traffic endpoint for a retention agent to pull an explanation "
     "for a specific disputed or reviewed customer, rather than adding "
     "explanation latency to every scoring call."),
])

add_heading("13. Immutable Prediction Logging", level=2)
add_para(
    "This is the direct answer to \"how are model predictions and the "
    "decisions based on them logged so they cannot be altered after the "
    "fact.\" A normal application log can be edited or deleted by anything "
    "with write access, including automated processes -- which is a "
    "governance gap for a model that can influence a customer-facing "
    "decision. If a prediction is later disputed by a customer, a "
    "regulator, or an internal audit, there needs to be a record that "
    "provably was not altered afterward."
)
add_image("../assets/governance_logging_diagram.png", width_in=6.5)
add_bullets([
    ("Azure Confidential Ledger ", "is the Azure-native service for this: "
     "entries are cryptographically hashed and chained, and the service "
     "runs inside hardware-backed confidential compute, so tampering is "
     "detectable rather than merely discouraged by access controls. This is "
     "the direct, real Azure equivalent of the \"blockchain-type "
     "immutability\" requirement -- a managed ledger service, not a "
     "custom-built blockchain, which is the appropriate scope for this."),
    ("Azure Key Vault ", "holds the signing key used to sign each "
     "prediction event before it is written, a separate permission boundary "
     "from the scoring service's own identity, so the signature cannot be "
     "forged by anything that only has scoring-service access."),
    ("What is logged: ", "a SHA-256 hash of the input (not the raw record, "
     "so the ledger does not become a second copy of customer PII needing "
     "its own retention policy), the model name and version, the "
     "prediction and probability, and a UTC timestamp. The raw input record "
     "itself is written separately to the scored-data store used for drift "
     "monitoring, under normal data-retention policy -- the ledger is the "
     "tamper-evident proof a prediction happened and what it was, not the "
     "system of record for customer data."),
    ("Fail-closed: ", "if the ledger write fails, the scoring request fails "
     "(Section 11) -- there is no code path that returns a prediction "
     "without also recording it."),
])
add_para(
    "Full design rationale: deployment/governance_logging.md. This "
    "specific design covers integrity and non-repudiation of the "
    "prediction record. It does not, on its own, cover bias/fairness "
    "testing (a separate pre-deployment gate, above) or the downstream "
    "retention action taken on a prediction -- that action should be "
    "logged by the application/UI layer with the same immutable pattern if "
    "the full decision chain is to be auditable end to end.",
    italic=True, size=10, color=MUTED_COLOR,
)
add_page_break()

add_heading("14. Continuous Improvement Strategy", level=2)
add_para(
    "How this model gets better over time, and where reinforcement "
    "learning genuinely fits versus where it does not:"
)
add_bullets([
    ("Champion/challenger retraining loop ", "(Section 10) -- scheduled "
     "weekly plus drift-triggered, with every candidate evaluated against "
     "the live production model before promotion. This is the primary "
     "mechanism by which the model improves as new data accumulates."),
    ("Active learning ", "on low-confidence predictions -- the "
     "needs_human_review flag from score.py (Section 11) naturally "
     "produces a queue of the cases the model is least certain about. "
     "Prioritizing these for outcome review and labeling improves the "
     "training set specifically where the model is weakest, rather than "
     "labeling uniformly at random."),
    ("Ground-truth feedback capture. ", "actual churn outcomes are joined "
     "back to historical predictions to compute the rolling recall/"
     "precision used for concept-drift detection (Section 7) -- this "
     "feedback loop is what makes drift detection possible at all, not an "
     "optional add-on."),
])
add_para(
    "On reinforcement learning specifically, stated directly rather than "
    "implied: churn prediction is a supervised classification problem, and "
    "reinforcement learning is not the right tool for it -- there is no "
    "sequential decision process to optimize inside \"will this customer "
    "churn.\" Where RL, specifically contextual bandits, genuinely does fit "
    "is one level downstream: once a customer is flagged at-risk, which "
    "retention offer to send is a repeated decision with observable "
    "reward (did the offer work), which is exactly the contextual-bandit "
    "setting. This is not built in this deliverable because the dataset "
    "contains no historical record of retention offers made or their "
    "outcomes -- there is nothing to learn a policy from yet. It is scoped "
    "here as the correct next step once that data starts to exist, not "
    "retrofitted onto the churn classifier where it does not belong."
)
add_para(
    "On generative AI specifically, also worth stating directly: an LLM is "
    "not the right tool for the core prediction problem here either -- this "
    "is structured tabular classification, and XGBoost/Random Forest/"
    "Logistic Regression are the appropriate tools for it, not a language "
    "model. Where an LLM could add real value is a layer above the "
    "analytics, not inside it: taking the SHAP-driven decision card from "
    "Section 6 and the customer risk profile, and generating a natural-"
    "language summary and suggested talking points for a retention agent, "
    "via Azure AI Foundry sitting on top of the SHAP output rather than "
    "replacing it. That is a scoped future extension, not part of this "
    "deliverable."
)

add_heading("15. Marketing Mix Model -- Applicability Assessment", level=2)
add_para(
    "Addressed directly rather than force-fitted: a Marketing Mix Model "
    "(MMM) was considered for this engagement and is not applicable to "
    "this dataset."
)
add_bullets([
    ("What MMM needs: ", "aggregated marketing spend by channel over time, "
     "regressed against a revenue or conversion KPI, typically at a "
     "weekly or monthly granularity across many campaign cycles."),
    ("What this dataset has: ", "customer-level, cross-sectional account "
     "and service data with no spend, channel, or campaign fields, and no "
     "time series at all -- a single snapshot. Building an MMM on it would "
     "not produce a fitted model with a fabricated result; it would produce "
     "a fabricated result presented as if it were fitted."),
])
add_para(
    "The applicable analogue for what MMM is generally used for -- \"where "
    "does spend actually move the outcome\" -- in a retention context is "
    "uplift modeling: estimating the incremental effect of a specific "
    "retention action on a specific customer, using a T-learner or causal "
    "forest trained against historical treatment/outcome data. This "
    "requires the same kind of data MMM would need (recorded retention "
    "actions and their outcomes), which does not yet exist for this "
    "dataset either. It is scoped as future work in Section 8 for the same "
    "reason RL-based offer selection is scoped as future work in Section 14 "
    "-- both are the correct tool for a question this specific dataset "
    "cannot yet answer, and both should be built once the underlying data "
    "exists rather than approximated now."
)

add_heading("16. Azure Platform Capabilities Referenced", level=2)
add_table(
    headers=["Capability", "Role in this architecture"],
    rows=[
        ["Azure ML Managed Feature Store", "Shared feature definitions between training and serving, preventing train/serve skew"],
        ["Azure ML Managed Online Endpoints", "Serving layer with built-in canary/traffic-split support used for safe rollout"],
        ["Azure ML Pipelines + Model Registry", "Training, evaluation, and champion/challenger versioning"],
        ["Azure Automated ML (AutoML)", "Available as a faster baseline-search step inside the training pipeline for future feature sets; not required here since the three-model comparison already covers baseline, bagging, and boosting"],
        ["Azure AI Foundry", "Current unified Azure AI platform surface; relevant if this model is later composed with a generative or agentic retention workflow"],
        ["Azure Confidential Ledger", "Immutable, tamper-evident prediction audit trail (Section 13)"],
        ["Azure Key Vault", "Signing/encryption keys for audit events, held in a separate permission boundary from the scoring service"],
        ["Azure Monitor / Application Insights", "Operational monitoring feeding the drift-detection job and canary health checks"],
        ["Azure DevOps Environments", "Manual approval gate for production promotion (Section 10)"],
    ],
    col_widths_cm=[5.2, 11.8],
    size=9,
)

add_heading("17. Deliverable Manifest", level=2)
add_table(
    headers=["Path", "Contents"],
    rows=[
        ["notebooks/telco_churn_analysis.ipynb", "Full executed analysis: EDA, feature engineering, 3 models, evaluation, lift/risk-segmentation, SHAP interpretation, monitoring plan"],
        ["docs/Telco_Churn_Writeup_and_Architecture.docx", "This document"],
        ["README.md", "Setup and run instructions for every deliverable in this repository"],
        ["app.py", "Streamlit decision-support prototype (Section 6)"],
        ["src/churn_features.py", "Shared cleaning/feature-engineering module used by app.py"],
        ["models/churn_xgboost_pipeline.joblib, model_metadata.json", "Persisted recommended model and its metadata, produced by the notebook"],
        ["deployment/training_pipeline.yml", "Azure ML CLI v2 training pipeline definition"],
        ["deployment/managed_online_endpoint.yml, managed_online_deployment.yml", "Serving endpoint and deployment definitions"],
        ["deployment/score.py", "Scoring entry script with guardrails and audit logging"],
        ["deployment/scoring_env.yml", "Pinned conda environment for the serving container"],
        ["deployment/monitoring_job.yml", "Scheduled drift/performance monitoring job"],
        ["deployment/azure-pipelines.yml", "CI/CD pipeline: CI, train/evaluate, approval gate, canary, full rollout"],
        ["deployment/governance_logging.md", "Immutable prediction logging design rationale"],
        ["assets/*.png", "Architecture, MLOps, and governance diagrams"],
        ["requirements.txt", "Pinned Python dependencies for local reproduction"],
        ["data/WA_Fn-UseC_-Telco-Customer-Churn.csv", "Source dataset"],
    ],
    col_widths_cm=[7.5, 9.5],
    size=8.5,
)

doc.save("../docs/Telco_Churn_Writeup_and_Architecture.docx")
print("Document written.")
